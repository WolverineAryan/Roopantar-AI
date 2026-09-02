import os
from pathlib import Path
from typing import Tuple
from pypdf import PdfReader
from docx import Document
import logging

logger = logging.getLogger(__name__)

def parse_txt(file_path: Path) -> str:
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read().strip()
    except Exception as e:
        logger.error(f"Error reading TXT file {file_path}: {e}")
        raise ValueError(f"Failed to read text file: {str(e)}")

def parse_pdf(file_path: Path) -> str:
    try:
        reader = PdfReader(str(file_path))
        extracted_text = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                extracted_text.append(page_text.strip())
        full_text = "\n\n".join(extracted_text)
        if not full_text.strip():
            raise ValueError("PDF contains no extractable text. It may be a scanned image.")
        return full_text
    except Exception as e:
        logger.error(f"Error reading PDF file {file_path}: {e}")
        raise ValueError(f"Failed to parse PDF document: {str(e)}")

def parse_docx(file_path: Path) -> str:
    try:
        doc = Document(str(file_path))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    paragraphs.append(" | ".join(row_text))
        return "\n\n".join(paragraphs)
    except Exception as e:
        logger.error(f"Error reading DOCX file {file_path}: {e}")
        raise ValueError(f"Failed to parse DOCX document: {str(e)}")

def parse_document(file_path: Path, filename: str) -> Tuple[str, str]:
    ext = Path(filename).suffix.lower()
    
    if ext in [".txt", ".md", ".log", ".csv", ".json"]:
        return parse_txt(file_path), "text"
    elif ext == ".pdf":
        return parse_pdf(file_path), "pdf"
    elif ext in [".docx", ".doc"]:
        return parse_docx(file_path), "docx"
    else:
        raise ValueError(f"Unsupported document format: {ext}")
