from pathlib import Path
from typing import Dict, Any
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def export_advisory_or_summary_to_docx(format_id: str, content: Dict[str, Any], output_path: Path) -> Path:
    doc = Document()
    
    # Page setup - 1 inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    title = content.get("title", content.get("infographic_title", content.get("video_title", "Roopantar-AI Deliverable")))
    
    # Title Heading
    h1 = doc.add_heading(level=0)
    run_title = h1.add_run(title)
    run_title.font.name = "Calibri"
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(15, 23, 42)
    
    if format_id == "advisory":
        # Advisory Metadata Header Box
        severity = content.get("severity", "Medium")
        adv_id = content.get("advisory_id", "ADV-2026-001")
        date_issued = content.get("date_issued", "2026-09-02")
        
        meta_p = doc.add_paragraph()
        r_id = meta_p.add_run(f"ADVISORY ID: {adv_id}  |  DATE: {date_issued}  |  SEVERITY: {severity.upper()}\n")
        r_id.font.bold = True
        r_id.font.size = Pt(11)
        if severity.lower() in ["high", "critical"]:
            r_id.font.color.rgb = RGBColor(220, 38, 38)
        else:
            r_id.font.color.rgb = RGBColor(37, 99, 235)
            
        # Summary
        doc.add_heading("1. Executive Overview", level=1)
        doc.add_paragraph(content.get("summary", ""))
        
        # Threat / Issue Breakdown
        doc.add_heading("2. Threat & Technical Analysis", level=1)
        for item in content.get("threat_or_issue_breakdown", []):
            p = doc.add_paragraph()
            r_h = p.add_run(f"• {item.get('heading', 'Finding')}: ")
            r_h.bold = True
            p.add_run(item.get("details", ""))
            
        # Impact
        doc.add_heading("3. Operational Impact Assessment", level=1)
        doc.add_paragraph(content.get("impact_assessment", ""))
        
        # Recommended Actions
        doc.add_heading("4. Recommended Mitigations & Action Items", level=1)
        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Priority"
        hdr_cells[1].text = "Target Team"
        hdr_cells[2].text = "Action Item"
        
        for action in content.get("recommended_actions", []):
            row_cells = table.add_row().cells
            row_cells[0].text = action.get("priority", "Immediate")
            row_cells[1].text = action.get("target_team", "Operations")
            row_cells[2].text = action.get("action", "")
            
        # References
        if content.get("references"):
            doc.add_heading("5. References & Standards", level=1)
            for ref in content.get("references", []):
                doc.add_paragraph(f"[{ref}]")

    elif format_id == "executive_summary":
        # Executive Summary docx
        doc.add_heading("1. Strategic Context", level=1)
        doc.add_paragraph(content.get("strategic_context", ""))
        
        doc.add_heading("2. Bottom Line Up Front (BLUF)", level=1)
        bluf_p = doc.add_paragraph()
        r_bluf = bluf_p.add_run(content.get("bottom_line_up_front", ""))
        r_bluf.bold = True
        
        doc.add_heading("3. Key Findings & Mission Impact", level=1)
        for finding in content.get("key_findings", []):
            p = doc.add_paragraph()
            p.add_run(f"• Area: {finding.get('area', '')}\n").bold = True
            p.add_run(f"  Observation: {finding.get('observation', '')}\n")
            p.add_run(f"  Impact: {finding.get('business_or_mission_impact', '')}\n")
            
        doc.add_heading("4. Decisions Required", level=1)
        for dec in content.get("decision_and_action_requirements", []):
            p = doc.add_paragraph()
            p.add_run(f"• Decision: {dec.get('decision_needed', '')} (Owner: {dec.get('stakeholder', '')} | Timeline: {dec.get('timeline', '')})")
            
        doc.add_heading("5. Resource Implications", level=1)
        doc.add_paragraph(content.get("resource_implications", ""))

    elif format_id == "video_package":
        # Video script docx
        doc.add_heading(f"Format: {content.get('target_format', '16:9')} | Duration: {content.get('target_duration', '90s')}", level=2)
        doc.add_paragraph(f"Logline: {content.get('logline', '')}").italic = True
        
        doc.add_heading("Scene-by-Scene Script & Storyboard", level=1)
        for scene in content.get("scenes", []):
            p = doc.add_paragraph()
            p.add_run(f"Scene {scene.get('scene_number')}: {scene.get('scene_name')} [{scene.get('timestamp_marker', '')}]\n").bold = True
            p.add_run(f"Visual: {scene.get('visual_description', '')}\n")
            p.add_run(f"Voiceover: \"{scene.get('narration_voiceover', '')}\"\n")
            p.add_run(f"On-Screen: {scene.get('on_screen_text', '')}\n")
            p.add_run(f"Audio Mood: {scene.get('audio_mood', '')}\n")

    else:
        # Generic document
        doc.add_heading("Generated Output Overview", level=1)
        for k, v in content.items():
            doc.add_heading(k.replace('_', ' ').title(), level=2)
            doc.add_paragraph(str(v))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
